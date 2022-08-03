import docx

# d = docx.Document('E:\Python Udemy\Automate boting stuff with Python programming course\WORD modifications', 'rb')

# p = d.paragraph[0].text
# p.runs[0].text
# p.runs[1].text

# p.runs[1].bold

# p.runs[1].underline = True

# p.runs[0].bold == None

# p.runs[3].text = 'italic and underlined'
# p.style = 'Title'

# d.save('E:\Python Udemy\Automate boting stuff with Python programming course\WORD modifications\\demoo.docx')
# ##########################################################################################################


# d = docx.Document()
# d.add_paragraph('Hello this is a paragraph')

# d.save('E:\Python Udemy\Automate boting stuff with Python programming course\WORD modifications\\demo4o.docx')

###################################

def getText(filename):
    doc = docx.Document(filename)
    fullText= []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('E:\\Python Udemy\\Automate boting stuff with Python programming course\\WORD modifications\\demoo34.docx'))